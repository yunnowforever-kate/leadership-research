"""DOCX 작성 에이전트: python-docx로 최종 Word 보고서 생성"""
import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
OUTPUT_DIR = os.environ.get("RESEARCH_OUTPUT_DIR", os.path.join(BASE_DIR, "outputs"))
KOREAN_FONT = "맑은 고딕"
ENGLISH_FONT = "Calibri"


# ── 헬퍼 ─────────────────────────────────────────────────────────────────────

def _set_font(run, size_pt: float, bold: bool = False, color: tuple = None):
    run.font.name = KOREAN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_paragraph_spacing(para, before_pt: float = 0, after_pt: float = 6, line_spacing: float = 1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = Pt(line_spacing * 12)


def _add_heading(doc: Document, text: str, level: int):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.name = KOREAN_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    _set_paragraph_spacing(para, before_pt=12 if level == 1 else 8, after_pt=4)
    return para


def _add_body(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_font(run, 10.5)
    _set_paragraph_spacing(para, before_pt=0, after_pt=4, line_spacing=1.5)
    return para


def _markdown_to_paragraphs(doc: Document, md_text: str):
    """마크다운 텍스트를 docx 단락으로 변환"""
    lines = md_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:], level=3)
        elif stripped.startswith("## "):
            _add_heading(doc, stripped[3:], level=2)
        elif stripped.startswith("# "):
            _add_heading(doc, stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ", "· ")):
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(stripped[2:])
            _set_font(run, 10.5)
            _set_paragraph_spacing(para, after_pt=2, line_spacing=1.4)
        elif re.match(r"^\d+\.", stripped):
            para = doc.add_paragraph(style="List Number")
            run = para.add_run(re.sub(r"^\d+\.\s*", "", stripped))
            _set_font(run, 10.5)
            _set_paragraph_spacing(para, after_pt=2, line_spacing=1.4)
        else:
            _add_body(doc, stripped)


def _add_horizontal_rule(doc: Document):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def _add_source_table(doc: Document, sources: dict):
    """참고문헌 테이블 생성"""
    _add_heading(doc, "참고문헌", level=1)

    papers = sources.get("papers", [])
    if papers:
        _add_heading(doc, "학술 논문", level=2)
        for i, p in enumerate(papers, 1):
            authors = ", ".join(p.get("authors", [])[:3])
            if len(p.get("authors", [])) > 3:
                authors += " 외"
            year = p.get("year", "")
            title = p.get("title", "")
            journal = p.get("journal", "")
            doi = p.get("doi_url", "")
            citations = p.get("citations", 0)
            line = f"[{i}] {authors} ({year}). {title}. {journal}. 인용:{citations}회. {doi}"
            _add_body(doc, line.strip(". "))

    hbr = sources.get("hbr_articles", [])
    if hbr:
        _add_heading(doc, "HBR 아티클", level=2)
        for i, a in enumerate(hbr, 1):
            title = a.get("title", "")
            author = a.get("author", "") or "HBR Staff"
            date = a.get("date", "")
            url = a.get("url", "")
            line = f"[{i}] {author} ({date}). {title}. Harvard Business Review. {url}"
            _add_body(doc, line.strip(". "))

    courses = sources.get("courses", [])
    if courses:
        _add_heading(doc, "온라인 강의", level=2)
        for i, c in enumerate(courses, 1):
            title = c.get("title", "")
            univ = c.get("university", "")
            platform = c.get("platform", "")
            instructor = c.get("instructor", "")
            url = c.get("url", "")
            free = "무료" if c.get("free") else "유료"
            line = f"[{i}] {instructor} ({univ}). {title}. {platform} [{free}]. {url}"
            _add_body(doc, line.strip(". "))


# ── 메인 ─────────────────────────────────────────────────────────────────────

def run(synthesis: dict) -> str:
    topic = synthesis.get("topic", "리서치")
    print(f"[DOCX 작성 에이전트] '{topic}' 문서 생성 시작...")
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(topic)
    _set_font(title_run, 22, bold=True, color=(31, 73, 125))
    title_para.paragraph_format.space_after = Pt(8)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("종합 리서치 보고서")
    _set_font(sub_run, 16, color=(68, 114, 196))

    doc.add_paragraph()
    _add_horizontal_rule(doc)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today_str = datetime.now().strftime("%Y년 %m월 %d일")
    meta_run = meta_para.add_run(f"작성일: {today_str}  |  자동 생성 보고서 v1.0")
    _set_font(meta_run, 10, color=(128, 128, 128))

    doc.add_page_break()

    # ── 목차 (섹션 제목 기반으로 동적 생성) ─────────────────────────────────────
    sections = synthesis.get("sections", {})
    section_order = ["의미", "방법", "사례", "학습자료"]

    _add_heading(doc, "목차", level=1)
    for i, key in enumerate(section_order, 1):
        if key not in sections:
            continue
        sec_title = sections[key].get("title", key)
        para = doc.add_paragraph()
        run = para.add_run(f"{i}. {sec_title}")
        _set_font(run, 10.5, bold=True)
        para.paragraph_format.space_after = Pt(3)
    para = doc.add_paragraph()
    run = para.add_run("참고문헌")
    _set_font(run, 10.5, bold=True)
    para.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ── 4개 섹션 ───────────────────────────────────────────────────────────────
    section_order = ["의미", "방법", "사례", "학습자료"]
    sections = synthesis.get("sections", {})

    for key in section_order:
        if key not in sections:
            continue
        sec = sections[key]
        _add_heading(doc, sec.get("title", key), level=1)
        content = sec.get("content", "")
        _markdown_to_paragraphs(doc, content)
        _add_horizontal_rule(doc)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 참고문헌 ───────────────────────────────────────────────────────────────
    _add_source_table(doc, synthesis.get("sources", {}))

    # ── 저장 ──────────────────────────────────────────────────────────────────
    today_file = datetime.now().strftime("%Y%m%d_%H%M")
    safe_topic = re.sub(r'[\\/*?:"<>|]', "_", topic)[:30]
    out_path = os.path.join(OUTPUT_DIR, f"{safe_topic}_{today_file}.docx")
    doc.save(out_path)

    print(f"[DOCX 작성 에이전트] 저장 완료 → {out_path}")
    return out_path


if __name__ == "__main__":
    import json, sys
    data_dir = os.path.join(BASE_DIR, "data")
    syn_path = os.path.join(data_dir, "synthesis.json")
    if not os.path.exists(syn_path):
        print("synthesis.json 없음. synthesis_agent.py 먼저 실행하세요.")
        sys.exit(1)
    with open(syn_path, encoding="utf-8") as f:
        synthesis = json.load(f)
    run(synthesis)
